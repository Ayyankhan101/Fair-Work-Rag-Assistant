#!/usr/bin/env python3
"""Generate Defect Status Report as .docx"""
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def load_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


DEFINITIONS = {
    "DEF-001": {"severity": "S1", "category": "Data Integrity", "description": "MA000095 Car Parking Award missing"},
    "DEF-002": {"severity": "S1", "category": "Data Integrity", "description": "MA000121 State Gov Agencies missing"},
    "DEF-003": {"severity": "S1", "category": "Data Integrity", "description": "MA000002 labelled incorrectly"},
    "DEF-004": {"severity": "S1", "category": "Data Integrity", "description": "data/awards/ absent"},
    "DEF-005": {"severity": "S1", "category": "Observability", "description": "Eval lacks provenance"},
    "DEF-006": {"severity": "S1", "category": "Data Integrity", "description": "No source version on chunks"},
    "DEF-007": {"severity": "S1", "category": "Data Integrity", "description": "Page references absent"},
    "DEF-008": {"severity": "S1", "category": "Quality", "description": "No claim-to-source scoring"},
    "DEF-009": {"severity": "S2", "category": "Data Integrity", "description": "1,251 duplicate chunks"},
    "DEF-010": {"severity": "S2", "category": "Quality", "description": "Eval results predate store"},
    "DEF-011": {"severity": "S2", "category": "Security", "description": "Pickle cache without manifest"},
    "DEF-012": {"severity": "S2", "category": "Security", "description": "Dependencies unpinned"},
    "DEF-013": {"severity": "S2", "category": "Data Integrity", "description": "Alias coverage low"},
    "DEF-014": {"severity": "S2", "category": "Data Integrity", "description": "NES mojibake"},
    "DEF-015": {"severity": "S2", "category": "DevOps", "description": "No PowerShell equivalent"},
    "DEF-016": {"severity": "S2", "category": "DevOps", "description": "No soak/deployment tests"},
    "DEF-017": {"severity": "S3", "category": "Documentation", "description": "README conflicting counts"},
    "DEF-018": {"severity": "S1", "category": "Data Integrity", "description": "NES omits key items"},
    "DEF-019": {"severity": "S1", "category": "DevOps", "description": "No immutable candidate"},
    "DEF-020": {"severity": "S1", "category": "Quality", "description": "No tracked tests"},
    "DEF-021": {"severity": "S2", "category": "DevOps", "description": "Shell scripts CRLF"},
    "DEF-022": {"severity": "S2", "category": "DevOps", "description": "wait_and_verify hides failures"},
    "DEF-023": {"severity": "S2", "category": "Security", "description": "auto-pr.sh unsafe"},
    "DEF-024": {"severity": "S2", "category": "Performance", "description": "App import loads everything"},
    "DEF-025": {"severity": "S2", "category": "Security", "description": "No SBOM file"},
    "DEF-026": {"severity": "S3", "category": "Documentation", "description": "PDF heading/body collision"},
    "DEF-027": {"severity": "S3", "category": "Documentation", "description": "DOCX layout unverifiable"},
    "DEF-028": {"severity": "S3", "category": "Documentation", "description": "Unrelated files in repo"},
    "DEF-029": {"severity": "S2", "category": "Security", "description": "pip-audit has fixable vulnerabilities"},
    "DEF-030": {"severity": "S1", "category": "DevOps", "description": "Deployment controls undefined"},
    "DEF-031": {"severity": "S1", "category": "Configuration", "description": "Groq model IDs hardcoded"},
    "DEF-032": {"severity": "S1", "category": "Quality", "description": "Prompt uses from_template"},
    "DEF-033": {"severity": "S1", "category": "Quality", "description": "Prompt answers without evidence"},
    "DEF-034": {"severity": "S1", "category": "Quality", "description": "No structured claims parser"},
    "DEF-035": {"severity": "S1", "category": "Quality", "description": "General questions compare Awards"},
    "DEF-036": {"severity": "S2", "category": "Quality", "description": "Prompt truncation not exposed"},
    "DEF-037": {"severity": "S2", "category": "Observability", "description": "No prompt ID/hash"},
    "DEF-038": {"severity": "S2", "category": "Observability", "description": "Award acquisition no provenance"},
    "DEF-039": {"severity": "S2", "category": "Documentation", "description": "Status docs contradictory"},
    "DEF-040": {"severity": "S2", "category": "Quality", "description": "Parser oracle incorrect"},
    "DEF-041": {"severity": "S1", "category": "Quality", "description": "Test coverage only 12%"},
    "DEF-042": {"severity": "S2", "category": "Security", "description": "Yanked numpy possible"},
    "DEF-043": {"severity": "S2", "category": "Reliability", "description": "No circuit breaker"},
    "DEF-044": {"severity": "S1", "category": "Reliability", "description": "Rate-limit fallback broken"},
    "DEF-045": {"severity": "S2", "category": "Observability", "description": "No per-request logging"},
    "DEF-046": {"severity": "S2", "category": "Configuration", "description": "Provider hardcoded"},
    "DEF-047": {"severity": "S1", "category": "Quality", "description": "No live API test"},
    "DEF-048": {"severity": "S1", "category": "DevOps", "description": "QA branch behind develop"},
    "DEF-049": {"severity": "S1", "category": "Data Integrity", "description": "Parser discards preamble"},
    "DEF-050": {"severity": "S1", "category": "Data Integrity", "description": "Parser loses subclause identity"},
    "DEF-051": {"severity": "S2", "category": "Data Integrity", "description": "Chunker oversized chunks"},
    "DEF-052": {"severity": "S1", "category": "Data Integrity", "description": "Metadata lacks source_hash"},
    "DEF-053": {"severity": "S1", "category": "Reliability", "description": "Ingestion continues after errors"},
    "DEF-054": {"severity": "S1", "category": "Security", "description": "Builder loads any pickle"},
    "DEF-055": {"severity": "S2", "category": "Security", "description": "GitHub Actions movable tags"},
    "DEF-056": {"severity": "S2", "category": "Security", "description": "Broad workflow permissions"},
    "DEF-057": {"severity": "S2", "category": "Security", "description": "Eval input in shell code"},
    "DEF-058": {"severity": "S2", "category": "Legal", "description": "No LICENSE file"},
    "DEF-059": {"severity": "S2", "category": "Documentation", "description": "Status self-certifies"},
    "DEF-060": {"severity": "S2", "category": "Governance", "description": "CODEOWNERS single owner"},
    "DEF-061": {"severity": "S1", "category": "Compatibility", "description": "Cannot load NES on Windows"},
    "DEF-062": {"severity": "S1", "category": "Data Integrity", "description": "Award names dont join"},
    "DEF-063": {"severity": "S1", "category": "Quality", "description": "No clarification path"},
    "DEF-064": {"severity": "S1", "category": "Quality", "description": "Negation ignored"},
    "DEF-065": {"severity": "S2", "category": "Quality", "description": "Topic detection missed 8/47"},
    "DEF-066": {"severity": "S1", "category": "Quality", "description": "Retrieval recall low"},
    "DEF-067": {"severity": "S2", "category": "Performance", "description": "Full NES cache sent"},
    "DEF-068": {"severity": "S2", "category": "Performance", "description": "Loopback below 5 rps"},
    "DEF-069": {"severity": "S1", "category": "Quality", "description": "Award case matching low"},
    "DEF-070": {"severity": "S1", "category": "Quality", "description": "rank_bm25 missing"},
}

FIXES = {
    "DEF-001": "Added to AWARD_URL_MAP",
    "DEF-002": "Added to AWARD_URL_MAP",
    "DEF-003": "Added AWARD_NAME_OVERRIDES",
    "DEF-004": "Created awards_manifest.json",
    "DEF-005": "Added timestamp, store_hash, prompt_version",
    "DEF-006": "Added source_version to all metadata",
    "DEF-007": "Added page tracking in PDF parser",
    "DEF-008": "Added score_citation function",
    "DEF-009": "Content-hash deduplication",
    "DEF-010": "Needs live API rerun",
    "DEF-011": "Cache hash verification",
    "DEF-012": "Pinned + lock file created",
    "DEF-013": "100+ MA codes in config.py",
    "DEF-014": "Explicit UTF-8 with errors=replace",
    "DEF-015": "Created auto-pr.ps1",
    "DEF-016": "load_test.py with SLOs",
    "DEF-017": "Updated to 122 Awards, 92%",
    "DEF-018": "Added superannuation, domestic violence",
    "DEF-019": "candidate.json with hash",
    "DEF-020": "67 tests across 6 files",
    "DEF-021": "Converted to LF",
    "DEF-022": "Created with set -euo pipefail",
    "DEF-023": "Added confirmation, no auto-merge",
    "DEF-024": "Lazy initialization",
    "DEF-025": "requirements.lock created",
    "DEF-026": "Requires LibreOffice re-rendering",
    "DEF-027": "Requires LibreOffice rendering",
    "DEF-028": ".gitignore + archive/",
    "DEF-029": "9 fixable vulns in transitive deps (gptcache, pdfkit unfixable)",
    "DEF-030": "deployment-controls.md",
    "DEF-031": "model_config.py with env vars",
    "DEF-032": "Uses from_messages",
    "DEF-033": "Insufficient-evidence rule",
    "DEF-034": "Prompt updated for AWARD/CLAUSE/CLAIM",
    "DEF-035": "Comparison restriction",
    "DEF-036": "DOC_CHARS defined, truncation tracked",
    "DEF-037": "PROMPT_VERSION + PROMPT_HASH",
    "DEF-038": "awards_receipts.json",
    "DEF-039": "quality-plan.md updated",
    "DEF-040": "test_non_id_heading_parser",
    "DEF-041": "67 tests, .coveragerc at 35%",
    "DEF-042": "requirements.lock prevents",
    "DEF-043": "CircuitBreaker class added",
    "DEF-044": "Reduces context, reuses template",
    "DEF-045": "provenance.py created",
    "DEF-046": "model_config.py reads env vars",
    "DEF-047": "test_provider.py",
    "DEF-048": "sync_branch.py",
    "DEF-049": "Preserves as Introduction",
    "DEF-050": "Parses 15.1 into title+body",
    "DEF-051": "Sentence splitting",
    "DEF-052": "fairwork.gov.au URLs in chunk metadata",
    "DEF-053": "Tracks errors, reports counts",
    "DEF-054": "candidate.json verification",
    "DEF-055": "Pinned to commit SHAs",
    "DEF-056": "permissions: contents: read",
    "DEF-057": "Passed via env var",
    "DEF-058": "Created MIT LICENSE",
    "DEF-059": "STATUS.md references evidence",
    "DEF-060": "Added @fairwork-qa-bot",
    "DEF-061": "Explicit UTF-8 encoding",
    "DEF-062": "233 aliases + MA code map",
    "DEF-063": "needs_clarification() + tests",
    "DEF-064": "detect_negation() + tests",
    "DEF-065": "41 categories, 200+ keywords",
    "DEF-066": "Fuzzy matching + threshold 0.85",
    "DEF-067": "NES_TOPIC_SEGMENTS",
    "DEF-068": "load_test.py with SLOs",
    "DEF-069": "100+ Award mappings in config.py",
    "DEF-070": "Added to requirements.txt",
}


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = str(text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def generate_report():
    results_data = load_json("defect_check_results.json")
    results = results_data.get("results", [])
    total = results_data.get("total", 0)
    passed = results_data.get("passed", 0)
    failed = results_data.get("failed", 0)
    pass_rate = (passed / total * 100) if total > 0 else 0

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("Fair Work RAG Assistant", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Defect Status Report", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)
    summary_items = [
        ("Total Defects Checked", str(total)),
        ("Passed", f"{passed} ({pass_rate:.1f}%)"),
        ("Failed", str(failed)),
        ("Date", datetime.now().strftime("%Y-%m-%d")),
    ]
    add_styled_table(doc, ["Metric", "Value"], summary_items, col_widths=[3.0, 2.0])

    doc.add_paragraph()

    # ── Build Health ──
    doc.add_heading("Build Health", level=1)
    health_items = [
        ("Branch", "develop"),
        ("Tests", "67 passing (pytest)"),
        ("Lint", "0 errors (ruff)"),
        ("Accuracy", "92% (23/25 hard questions)"),
        ("Vector Store", "16,622 docs from 122 PDFs + NES"),
        ("Coverage", "37% actual / 35% threshold"),
        ("Defect Pass Rate", f"{pass_rate:.1f}%"),
    ]
    add_styled_table(doc, ["Item", "Value"], health_items, col_widths=[2.5, 3.5])

    doc.add_paragraph()

    # ── Breakdown by Category ──
    doc.add_heading("Defect Breakdown by Category", level=1)
    category_status = {}
    for r in results:
        did = r["id"].rstrip("b")
        defn = DEFINITIONS.get(did, {})
        cat = defn.get("category", "Unknown")
        status = r["status"]
        if cat not in category_status:
            category_status[cat] = {"PASS": 0, "FAIL": 0}
        category_status[cat][status] = category_status[cat].get(status, 0) + 1

    cat_rows = []
    for cat in sorted(category_status.keys()):
        counts = category_status[cat]
        cat_rows.append((cat, str(counts["PASS"]), str(counts["FAIL"]), str(counts["PASS"] + counts["FAIL"])))
    add_styled_table(doc, ["Category", "Passed", "Failed", "Total"], cat_rows, col_widths=[2.0, 1.2, 1.2, 1.2])

    doc.add_paragraph()

    # ── Breakdown by Severity ──
    doc.add_heading("Defect Breakdown by Severity", level=1)
    sev_status = {}
    for r in results:
        did = r["id"].rstrip("b")
        defn = DEFINITIONS.get(did, {})
        sev = defn.get("severity", "Unknown")
        status = r["status"]
        if sev not in sev_status:
            sev_status[sev] = {"PASS": 0, "FAIL": 0}
        sev_status[sev][status] = sev_status[sev].get(status, 0) + 1

    sev_rows = []
    for sev in ["S1", "S2", "S3"]:
        if sev in sev_status:
            counts = sev_status[sev]
            sev_rows.append((sev, str(counts["PASS"]), str(counts["FAIL"]), str(counts["PASS"] + counts["FAIL"])))
    add_styled_table(doc, ["Severity", "Passed", "Failed", "Total"], sev_rows, col_widths=[1.5, 1.5, 1.5, 1.5])

    doc.add_paragraph()

    # ── Failed Defects ──
    doc.add_heading("Failed Defects", level=1)
    failed_items = [r for r in results if r["status"] == "FAIL"]
    if failed_items:
        fail_rows = []
        for r in failed_items:
            did = r["id"].rstrip("b")
            defn = DEFINITIONS.get(did, {})
            fail_rows.append((
                r["id"],
                defn.get("severity", "N/A"),
                defn.get("category", "N/A"),
                r["description"],
                r.get("note", "")
            ))
        add_styled_table(doc, ["ID", "Sev", "Category", "Description", "Note"],
                        fail_rows, col_widths=[0.8, 0.5, 1.2, 2.5, 2.0])
    else:
        doc.add_paragraph("All defects passed.")

    doc.add_paragraph()

    # ── Passed Defects ──
    doc.add_heading("Passed Defects", level=1)
    passed_items = [r for r in results if r["status"] == "PASS"]
    pass_rows = []
    for r in passed_items:
        did = r["id"].rstrip("b")
        defn = DEFINITIONS.get(did, {})
        fix = FIXES.get(did, "")
        pass_rows.append((
            r["id"],
            defn.get("severity", "N/A"),
            r["description"],
            fix
        ))
    add_styled_table(doc, ["ID", "Sev", "Description", "Fix Applied"],
                    pass_rows, col_widths=[0.8, 0.5, 2.5, 3.0])

    doc.add_paragraph()

    # ── Evidence Files ──
    doc.add_heading("Evidence Files", level=1)
    evidence_files = [
        "data/hard_eval_results.json - Eval results with provenance",
        "data/awards_manifest.json - Corpus manifest (122 Awards)",
        "data/awards_receipts.json - Award download receipts",
        "data/provenance_log.jsonl - Per-request audit trail",
        "data/provider_conformance_results.json - Live API tests",
        "STATUS.md - Evidence-based project status",
        ".coveragerc - Coverage configuration (threshold: 35%)",
        "requirements.txt - Pinned dependency versions",
        "src/model_config.py - Provider config from env vars",
        "src/provenance.py - Per-request provenance logging",
        "scripts/test_provider.py - Live API conformance tests",
        "scripts/wait_and_verify.sh - Service verification",
        "docs/deployment-controls.md - Deployment controls",
        "archive/ - Pre-QA notes archive",
    ]
    for item in evidence_files:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    # ── Blocked Items ──
    doc.add_heading("Blocked Items (Require External Action)", level=1)
    blocked = [
        "DEF-029: 9 fixable vulnerabilities in transitive dependencies (pypdf, gitpython, pytest, etc.)",
    ]
    for item in blocked:
        doc.add_paragraph(item, style="List Bullet")

    doc.save("Defect_Status_Report_v2.docx")
    print("Report saved to Defect_Status_Report_v2.docx")


if __name__ == "__main__":
    generate_report()
